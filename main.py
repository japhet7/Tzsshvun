#!/usr/bin/env python3
"""
TZSSHVUN V1 - Ethical Vulnerability Scanner
Author: Japhet Munisi
"""

import sys
import subprocess
import socket
import requests
from bs4 import BeautifulSoup
from datetime import datetime
from docx import Document
from docx.shared import Pt
import nmap

# -------------------------
# CONFIG (Author Info)
# -------------------------
AUTHOR = "Japhet Munisi"
EMAIL = "munisi.japhet@gmail.com"
WHATSAPP = "+255784121281"
TOOL_NAME = "TZSSHVUN V1"

# -------------------------
# Dependency Check
# -------------------------
def check_dependencies():
    try:
        import nmap
        import requests
        import bs4
        import docx
    except ImportError as e:
        print(f"[ERROR] Missing dependency: {e}")
        print("Run: pip install -r requirements.txt")
        sys.exit(1)

# -------------------------
# Network Scanner
# -------------------------
def network_scan(target):
    results = []
    try:
        scanner = nmap.PortScanner()
        print(f"[+] Scanning {target}...")

        scanner.scan(hosts=target, arguments='-sV -O --script vuln')

        for host in scanner.all_hosts():
            host_data = {
                "host": host,
                "state": scanner[host].state(),
                "protocols": []
            }

            for proto in scanner[host].all_protocols():
                ports = scanner[host][proto].keys()
                for port in ports:
                    service = scanner[host][proto][port]
                    entry = {
                        "port": port,
                        "state": service['state'],
                        "name": service.get('name'),
                        "product": service.get('product'),
                        "version": service.get('version'),
                        "scripts": service.get('script', {})
                    }
                    host_data["protocols"].append(entry)

            results.append(host_data)

    except Exception as e:
        print(f"[ERROR] Network scan failed: {e}")

    return results

# -------------------------
# Web Vulnerability Scanner
# -------------------------
def web_scan(url):
    findings = []

    try:
        response = requests.get(url, timeout=10)
        headers = response.headers

        # Security Headers Check
        required_headers = [
            "Content-Security-Policy",
            "X-Frame-Options",
            "Strict-Transport-Security",
            "X-Content-Type-Options"
        ]

        for header in required_headers:
            if header not in headers:
                findings.append({
                    "type": "Missing Security Header",
                    "detail": header
                })

        # Directory Listing Check
        if "Index of /" in response.text:
            findings.append({
                "type": "Open Directory",
                "detail": "Directory listing enabled"
            })

        # Parse Forms for potential XSS/SQLi
        soup = BeautifulSoup(response.text, "html.parser")
        forms = soup.find_all("form")

        for form in forms:
            inputs = form.find_all("input")
            for inp in inputs:
                name = inp.get("name")
                if name:
                    findings.append({
                        "type": "Potential Input Field",
                        "detail": f"Input '{name}' may be vulnerable to XSS/SQLi"
                    })

    except Exception as e:
        print(f"[ERROR] Web scan failed: {e}")

    return findings

# -------------------------
# Report Generator
# -------------------------
def generate_report(target, net_results, web_results):
    doc = Document()

    # Cover Page
    doc.add_heading(TOOL_NAME, 0)
    doc.add_paragraph("Professional Vulnerability Assessment Report")
    doc.add_paragraph(f"Target: {target}")
    doc.add_paragraph(f"Date: {datetime.now()}")
    doc.add_page_break()

    # Executive Summary
    doc.add_heading("Executive Summary", 1)
    doc.add_paragraph(
        "This report outlines identified vulnerabilities from automated scanning. "
        "Immediate remediation is recommended to reduce risk exposure."
    )

    # Technical Findings
    doc.add_heading("Technical Findings", 1)

    # Network Findings
    doc.add_heading("Network Scan Results", 2)
    for host in net_results:
        doc.add_paragraph(f"Host: {host['host']} ({host['state']})")
        for service in host["protocols"]:
            doc.add_paragraph(
                f"Port {service['port']} - {service['name']} "
                f"{service['product']} {service['version']}"
            )
            if service["scripts"]:
                doc.add_paragraph(f"Vulnerabilities: {service['scripts']}")

    # Web Findings
    doc.add_heading("Web Vulnerabilities", 2)
    for vuln in web_results:
        doc.add_paragraph(f"{vuln['type']} - {vuln['detail']}")

    # Recommendations Table
    doc.add_heading("Remediation & Recommendations", 1)
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Vulnerability"
    table.rows[0].cells[1].text = "Recommendation"

    for vuln in web_results:
        row = table.add_row().cells
        row[0].text = vuln['type']
        row[1].text = "Apply security best practices and patch system."

    # Author Info
    doc.add_heading("Author Information", 1)
    doc.add_paragraph(f"Author: {AUTHOR}")
    doc.add_paragraph(f"Email: {EMAIL}")
    doc.add_paragraph(f"WhatsApp: {WHATSAPP}")

    filename = f"TZSSHVUN_Report_{target}.docx"
    doc.save(filename)

    print(f"[+] Report saved: {filename}")

# -------------------------
# Main
# -------------------------
def main():
    check_dependencies()

    if len(sys.argv) < 2:
        print("Usage: python main.py <target_ip_or_url>")
        sys.exit(1)

    target = sys.argv[1]

    print("[*] Starting TZSSHVUN V1...\n")

    net_results = network_scan(target)

    web_results = []
    if target.startswith("http"):
        web_results = web_scan(target)

    generate_report(target, net_results, web_results)

    print("[+] Scan completed successfully.")

if __name__ == "__main__":
    main()            sys.exit(1)

    def scan(self):
        print(f"[*] Initializing deep scan on {self.target}...")
        nm = nmap.PortScanner()
        # Performs service versioning and runs default vulnerability scripts
        nm.scan(self.target, arguments='-sV --script=vuln')
        
        if self.target in nm.all_hosts():
            self.results = nm[self.target]
        else:
            print("[-] Target host unreachable.")
            sys.exit(1)

    def generate_report(self):
        filename = f"Report_{self.target.replace('.', '_')}.docx"
        doc = Document()

        # Title Section
        header = doc.add_heading(f'{self.tool_name} Assessment Report', 0)
        doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Target: {self.target}")
        doc.add_paragraph(f"Auditor: {self.author}")

        # Executive Summary
        doc.add_heading('1. Executive Summary', level=1)
        doc.add_paragraph(
            "This document provides a technical security assessment of the target host. "
            "The scan identified open ports, service versions, and potential vulnerabilities."
        )

        # Technical Evidence
        doc.add_heading('2. Technical Evidence', level=1)
        if 'tcp' in self.results:
            for port, info in self.results['tcp'].items():
                p = doc.add_paragraph()
                run = p.add_run(f"Port {port} | Service: {info['name']}")
                run.bold = True
                doc.add_paragraph(f"Version: {info['product']} {info['version']}")
                
                if 'script' in info:
                    doc.add_heading('Vulnerability Data Found:', level=2)
                    for script_name, output in info['script'].items():
                        # Create a "code block" look for evidence
                        table = doc.add_table(rows=1, cols=1)
                        table.style = 'Light List Accent 1'
                        cell = table.rows[0].cells[0]
                        cell.text = f"Script ID: {script_name}\n\n{output}"

        # Solutions
        doc.add_heading('3. Solutions & Recommendations', level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        headers = table.rows[0].cells
        headers[0].text = 'Issue Found'
        headers[1].text = 'Recommended Fix'
        headers[2].text = 'Severity'

        if 'tcp' in self.results:
            for port, info in self.results['tcp'].items():
                row = table.add_row().cells
                row[0].text = f"Open Port {port} ({info['name']})"
                row[1].text = f"Update {info['product']} to the latest version or close port if not needed."
                row[2].text = "High" if 'script' in info else "Low"

        doc.save(filename)
        print(f"[+] Success! Professional report saved as: {filename}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python main.py <target_ip>")
        sys.exit(1)
    
    target = sys.argv[1]
    engine = TZSSHVUN(target)
    engine.banner()
    engine.check_nmap()
    engine.scan()
    engine.generate_report()
